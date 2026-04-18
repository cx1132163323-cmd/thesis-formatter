"""
Hunan University Master Thesis Formatter v3

CLI:
    python format_thesis.py [input.docx [output.docx]] [--title "Full Thesis Title"]

Web usage:
    from format_thesis import run_formatter
"""

import argparse
import io
import os
import re
import sys

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

HERE = os.path.dirname(os.path.abspath(__file__))
DEFAULT_ODD_HEADER_TEXT = "Dissertation of Master Degree"
DEFAULT_FALLBACK_TITLE = "Dissertation of Master Degree"


def xml_el(tag, **attrs):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), str(v))
    return el


def set_snap_to_grid(pf, on=False):
    p_pr = pf._element
    el = p_pr.find(qn("w:snapToGrid"))
    if el is None:
        el = OxmlElement("w:snapToGrid")
        p_pr.insert(0, el)
    el.set(qn("w:val"), "1" if on else "0")


def set_para_fmt(
    pf,
    *,
    alignment=None,
    space_before=None,
    space_after=None,
    first_line_indent=None,
    left_indent=None,
    line_spacing=None,
    line_spacing_rule=None,
    page_break_before=None,
    snap_to_grid=False,
):
    if alignment is not None:
        pf.alignment = alignment
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if line_spacing_rule is not None:
        pf.line_spacing_rule = line_spacing_rule
    if page_break_before is not None:
        pf.page_break_before = page_break_before
    set_snap_to_grid(pf, snap_to_grid)


def set_style_font(font, *, name=None, size=None, bold=None, italic=None):
    if name is not None:
        font.name = name
        r_pr = font._element
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        r_fonts.set(qn("w:ascii"), name)
        r_fonts.set(qn("w:hAnsi"), name)
    if size is not None:
        font.size = size
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic


def border_el(tag, val, sz, color="000000"):
    el = OxmlElement(f"w:{tag}")
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)
    return el


def get_or_create_style(doc, name, base_name):
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base_name]
        return style


def para_visible_text(p_elem):
    return "".join((t.text or "") for t in p_elem.findall(".//" + qn("w:t")))


def is_heading1_elem(p_elem):
    p_style = p_elem.find(".//" + qn("w:pStyle"))
    if p_style is None:
        return False
    val = p_style.get(qn("w:val"), "").replace(" ", "").replace("-", "").lower()
    return val in ("heading1", "1")


def has_only_page_break(p_elem):
    for br in p_elem.findall(".//" + qn("w:br")):
        if br.get(qn("w:type")) == "page" and not para_visible_text(p_elem).strip():
            return True
    return False


def make_field_run(fld_char_type=None, instr_text=None):
    r = OxmlElement("w:r")
    if fld_char_type:
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), fld_char_type)
        if fld_char_type == "begin":
            fld_char.set(qn("w:dirty"), "1")
        r.append(fld_char)
    if instr_text:
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {instr_text} "
        r.append(instr)
    return r


def build_toc_p(field_instruction):
    p = OxmlElement("w:p")
    p.append(make_field_run(fld_char_type="begin"))
    p.append(make_field_run(instr_text=field_instruction))
    p.append(make_field_run(fld_char_type="separate"))
    placeholder_r = OxmlElement("w:r")
    placeholder_t = OxmlElement("w:t")
    placeholder_t.text = "[Open in Word -> Ctrl+A -> F9 to update]"
    placeholder_r.append(placeholder_t)
    p.append(placeholder_r)
    p.append(make_field_run(fld_char_type="end"))
    return p


def build_heading1_p(text, style_id="Heading1"):
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    p_style = OxmlElement("w:pStyle")
    p_style.set(qn("w:val"), style_id)
    p_pr.append(p_style)
    p.append(p_pr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def add_page_number_field(paragraph):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = paragraph._p

    def field_run(fld_char_type=None, instr=None, font="Times New Roman", sz_pt=10):
        r = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), font)
        r_fonts.set(qn("w:hAnsi"), font)
        r_pr.append(r_fonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(sz_pt * 2))
        r_pr.append(sz)
        r.append(r_pr)
        if fld_char_type:
            fld_char = OxmlElement("w:fldChar")
            fld_char.set(qn("w:fldCharType"), fld_char_type)
            r.append(fld_char)
        if instr:
            instr_text = OxmlElement("w:instrText")
            instr_text.set(qn("xml:space"), "preserve")
            instr_text.text = f" {instr} "
            r.append(instr_text)
        return r

    p.append(field_run(fld_char_type="begin"))
    p.append(field_run(instr="PAGE"))
    p.append(field_run(fld_char_type="end"))


def find_thesis_title(doc, override=None):
    if override:
        return override
    title = (doc.core_properties.title or "").strip()
    if len(title) > 15:
        return title
    for p in doc.paragraphs[:80]:
        txt = p.text.strip()
        if (
            25 < len(txt) < 200
            and txt[0].isupper()
            and re.search(r"[A-Za-z]{4,}", txt)
            and not re.match(
                r"(Dissertation|Candidate|Supervisor|Author|Department|Subject|Research|Date|Hunan|I solemnly)",
                txt,
            )
        ):
            return txt
    return DEFAULT_FALLBACK_TITLE


def in_list_section(para_text):
    return bool(re.search(r"\t\d+\s*$", para_text))


def is_centered_h1(text):
    centered_exact = {
        "Abstract",
        "摘要",
        "References",
        "Acknowledgement",
        "CONTENTS",
        "Contents",
    }
    centered_prefix = (
        "Appendix",
        "List of",
        "Table of",
        "LIST OF",
        "TABLE OF",
        "list of",
        "table of",
    )
    t = text.strip()
    return t in centered_exact or any(t.startswith(prefix) for prefix in centered_prefix)


def apply_three_line_table(table):
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)

    tbl_borders = OxmlElement("w:tblBorders")
    tbl_borders.append(border_el("top", "single", 12))
    tbl_borders.append(border_el("left", "none", 0))
    tbl_borders.append(border_el("bottom", "single", 12))
    tbl_borders.append(border_el("right", "none", 0))
    tbl_borders.append(border_el("insideH", "none", 0))
    tbl_borders.append(border_el("insideV", "none", 0))
    tbl_pr.append(tbl_borders)

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tc = cell._tc
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is None:
                tc_pr = OxmlElement("w:tcPr")
                tc.insert(0, tc_pr)
            for old in tc_pr.findall(qn("w:tcBorders")):
                tc_pr.remove(old)
            if row_idx == 0:
                tc_borders = OxmlElement("w:tcBorders")
                tc_borders.append(border_el("bottom", "single", 6))
                tc_pr.append(tc_borders)

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.paragraph_format.alignment is None:
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    if run.font.name is None:
                        run.font.name = "Times New Roman"
                    if run.font.size is None or run.font.size == Pt(12):
                        run.font.size = Pt(10.5)


def promote_embedded_list_heading(doc, keyword):
    for para in doc.paragraphs:
        if para.style.name != "Heading 1" and keyword.lower() in para.text.strip().lower():
            if not in_list_section(para.text):
                para.style = doc.styles["Heading 1"]
                return True
    return False


def replace_or_insert_toc(doc, heading_keyword, field_instr, insert_before_keyword=None):
    paras = list(doc.paragraphs)
    heading_p = None

    for p in paras:
        if p.style.name == "Heading 1" and heading_keyword.lower() in p.text.lower():
            heading_p = p
            break

    if heading_p is not None:
        collecting = False
        to_del = []
        for p in paras:
            if p is heading_p:
                collecting = True
                continue
            if collecting:
                if p.style.name == "Heading 1":
                    break
                to_del.append(p._element)
        for elem in to_del:
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)

        heading_p._element.addnext(build_toc_p(field_instr))
        return f"replaced content in '{heading_p.text.strip()}'"

    if insert_before_keyword:
        ref_p = None
        for p in paras:
            if p.style.name == "Heading 1" and insert_before_keyword.lower() in p.text.lower():
                ref_p = p
                break
        if ref_p is None:
            return f"NOT FOUND ('{heading_keyword}') and anchor '{insert_before_keyword}' also missing"

        toc_p_elem = build_toc_p(field_instr)
        heading_p_elem = build_heading1_p(heading_keyword.title())
        ref_p._element.addprevious(toc_p_elem)
        ref_p._element.addprevious(heading_p_elem)
        return f"inserted new '{heading_keyword.title()}' section before '{ref_p.text.strip()}'"

    return f"NOT FOUND ('{heading_keyword}')"


def process_document(doc, title_override=None, log_func=None):
    logs = []

    def log(message):
        logs.append(message)
        if log_func:
            log_func(message)

    thesis_title = find_thesis_title(doc, title_override)
    log(f"Thesis title: {thesis_title[:80]!r}")

    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)
        sec.header_distance = Cm(1.1)
        sec.footer_distance = Cm(1.1)
    log("Page setup: A4, top/bottom 2.2 cm, left/right 2.5 cm, header/footer 1.1 cm")

    normal = doc.styles["Normal"]
    set_style_font(normal.font, name="Times New Roman", size=Pt(12), bold=False)
    set_para_fmt(
        normal.paragraph_format,
        line_spacing=1.25,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        space_before=Pt(0),
        space_after=Pt(0),
        first_line_indent=Pt(24),
        snap_to_grid=False,
    )

    h1 = doc.styles["Heading 1"]
    set_style_font(h1.font, name="Times New Roman", size=Pt(18), bold=True)
    set_para_fmt(
        h1.paragraph_format,
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
        space_before=Pt(18),
        space_after=Pt(18),
        first_line_indent=Pt(0),
        left_indent=Pt(0),
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        page_break_before=True,
        snap_to_grid=False,
    )

    h2 = doc.styles["Heading 2"]
    set_style_font(h2.font, name="Times New Roman", size=Pt(15), bold=True)
    set_para_fmt(
        h2.paragraph_format,
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
        space_before=Pt(12),
        space_after=Pt(12),
        first_line_indent=Pt(0),
        left_indent=Pt(0),
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        snap_to_grid=False,
    )

    h3 = doc.styles["Heading 3"]
    set_style_font(h3.font, name="Times New Roman", size=Pt(14), bold=True)
    set_para_fmt(
        h3.paragraph_format,
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
        space_before=Pt(6),
        space_after=Pt(6),
        first_line_indent=Pt(0),
        left_indent=Pt(0),
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        snap_to_grid=False,
    )

    h4 = get_or_create_style(doc, "Heading 4", "Heading 3")
    set_style_font(h4.font, name="Times New Roman", size=Pt(12), bold=True)
    set_para_fmt(
        h4.paragraph_format,
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
        space_before=Pt(6),
        space_after=Pt(6),
        first_line_indent=Pt(0),
        left_indent=Pt(0),
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        snap_to_grid=False,
    )

    cap_style = doc.styles["Caption"]
    set_style_font(cap_style.font, name="Times New Roman", size=Pt(12), bold=False, italic=False)
    set_para_fmt(
        cap_style.paragraph_format,
        line_spacing=1.25,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        space_before=Pt(6),
        space_after=Pt(6),
        first_line_indent=Pt(0),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        snap_to_grid=False,
    )

    fig_cap = get_or_create_style(doc, "Figure Caption", "Caption")
    set_style_font(fig_cap.font, name="Times New Roman", size=Pt(12), bold=False, italic=False)
    set_para_fmt(
        fig_cap.paragraph_format,
        line_spacing=1.25,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        space_before=Pt(6),
        space_after=Pt(6),
        first_line_indent=Pt(0),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        snap_to_grid=False,
    )

    tab_cap = get_or_create_style(doc, "Table Caption", "Caption")
    set_style_font(tab_cap.font, name="Times New Roman", size=Pt(12), bold=False, italic=False)
    set_para_fmt(
        tab_cap.paragraph_format,
        line_spacing=1.25,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        space_before=Pt(6),
        space_after=Pt(6),
        first_line_indent=Pt(0),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        snap_to_grid=False,
    )
    log("Styles: Normal, Heading 1-4, Caption, Figure Caption, Table Caption")

    caption_styles = {"Caption", "Normal", "Normal (Web)", "Body Text"}
    caption_figure = re.compile(r"^(Fig\.?|Figure)\s+\d", re.I)
    caption_table = re.compile(r"^(Tab\.?|Table)\s+\d", re.I)
    fig_n = 0
    tab_n = 0
    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt or in_list_section(txt):
            continue
        if para.style.name in caption_styles:
            if caption_figure.match(txt):
                para.style = doc.styles["Figure Caption"]
                fig_n += 1
            elif caption_table.match(txt):
                para.style = doc.styles["Table Caption"]
                tab_n += 1
    log(f"Re-assigned captions: {fig_n} -> Figure Caption, {tab_n} -> Table Caption")

    left_n = 0
    ctr_n = 0
    for para in doc.paragraphs:
        if para.style.name != "Heading 1":
            continue
        pf = para.paragraph_format
        pf.first_line_indent = Pt(0)
        pf.left_indent = Pt(0)
        if is_centered_h1(para.text):
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ctr_n += 1
        else:
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            left_n += 1
    log(f"Heading 1 alignment: {left_n} left, {ctr_n} centered")

    ref_start_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1" and "Reference" in p.text:
            ref_start_idx = i
            break

    keep_sizes = {Pt(10.5), Pt(12), Pt(14), Pt(15), Pt(18), Pt(9), Pt(10)}
    for i, para in enumerate(doc.paragraphs):
        if para.style.name != "Normal":
            continue
        pf = para.paragraph_format
        if pf.space_before is not None:
            try:
                if pf.space_before.pt != 0:
                    pf.space_before = Pt(0)
            except Exception:
                pass
        if pf.space_after is not None:
            try:
                if pf.space_after.pt != 0:
                    pf.space_after = Pt(0)
            except Exception:
                pass
        if pf.line_spacing is not None:
            try:
                if abs(float(pf.line_spacing) - 1.25) > 0.05:
                    pf.line_spacing = 1.25
                    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            except Exception:
                pass
        for run in para.runs:
            if run.font.size is not None and run.font.size not in keep_sizes:
                run.font.size = None
        if ref_start_idx is not None and i > ref_start_idx and re.match(r"^\s*\[\d+\]", para.text):
            pf.first_line_indent = Pt(-18)
            pf.left_indent = Pt(18)
    log("Body text overrides normalised")

    first_content_h1_idx = None
    body_elem = doc.element.body
    elems = [(c, c.tag.split("}")[-1] == "p") for c in list(body_elem)]

    for idx, (elem, is_p) in enumerate(elems):
        if is_p and is_heading1_elem(elem) and para_visible_text(elem).strip():
            first_content_h1_idx = idx
            break

    to_remove = []
    blank_run = 0
    for idx, (elem, is_p) in enumerate(elems):
        if not is_p:
            blank_run = 0
            continue
        if has_only_page_break(elem):
            to_remove.append(elem)
            continue

        is_blank = not para_visible_text(elem).strip()
        if is_blank:
            in_body = first_content_h1_idx is not None and idx >= first_content_h1_idx
            if in_body:
                prev_elem = elems[idx - 1][0] if idx > 0 else None
                next_elem = elems[idx + 1][0] if idx < len(elems) - 1 else None
                prev_h1 = prev_elem is not None and elems[idx - 1][1] and is_heading1_elem(prev_elem)
                next_h1 = next_elem is not None and elems[idx + 1][1] and is_heading1_elem(next_elem)
                if prev_h1 or next_h1 or blank_run >= 1:
                    to_remove.append(elem)
                else:
                    blank_run += 1
            else:
                if blank_run >= 2:
                    to_remove.append(elem)
                else:
                    blank_run += 1
        else:
            blank_run = 0

    for elem, is_p in elems:
        if is_p and is_heading1_elem(elem) and not para_visible_text(elem).strip():
            to_remove.append(elem)

    for elem in to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
    log(f"Cleanup: removed {len(to_remove)} blank/page-break-only paragraphs")

    for table in doc.tables:
        apply_three_line_table(table)
    log(f"Three-line tables: {len(doc.tables)} tables updated")

    for kw in ("list of tables", "list of figures"):
        promote_embedded_list_heading(doc, kw)

    toc_tasks = [
        ("contents", r'TOC \o "1-3" \h \z \u', "list of figures"),
        ("list of figures", r'TOC \h \z \t "Figure Caption,1"', None),
        ("list of tables", r'TOC \h \z \t "Table Caption,1"', "references"),
    ]
    log("TOC field codes:")
    for keyword, field, anchor in toc_tasks:
        result = replace_or_insert_toc(doc, keyword, field, anchor)
        log(f"  {keyword}: {result}")
    log("  Open in Word -> Ctrl+A -> F9 to populate page numbers")

    for sec in doc.sections:
        footer = sec.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        add_page_number_field(fp)
    log("Footer: centered PAGE field")

    settings_el = doc.settings.element
    if settings_el.find(qn("w:evenAndOddHeaders")) is None:
        settings_el.append(OxmlElement("w:evenAndOddHeaders"))

    def configure_header(para, text, font_size_pt=9):
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(font_size_pt)

        p_pr = para._p.get_or_add_pPr()
        for old in p_pr.findall(qn("w:pBdr")):
            p_pr.remove(old)
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "thickThinSmallGap")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    even_header_text = thesis_title[:80]
    for sec in doc.sections:
        odd_para = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
        configure_header(odd_para, DEFAULT_ODD_HEADER_TEXT)

        even_para = (
            sec.even_page_header.paragraphs[0]
            if sec.even_page_header.paragraphs
            else sec.even_page_header.add_paragraph()
        )
        configure_header(even_para, even_header_text)
    log(f"Headers: odd = '{DEFAULT_ODD_HEADER_TEXT}', even = title")

    return thesis_title, logs


def run_formatter(input_source, *, title=None, output_path=None, log_func=None):
    if isinstance(input_source, (bytes, bytearray)):
        doc = Document(io.BytesIO(input_source))
    else:
        doc = Document(input_source)

    detected_title, logs = process_document(doc, title_override=title, log_func=log_func)

    buffer = io.BytesIO()
    doc.save(buffer)
    output_bytes = buffer.getvalue()

    saved_path = None
    if output_path:
        with open(output_path, "wb") as f:
            f.write(output_bytes)
        saved_path = output_path

    return {
        "output_bytes": output_bytes,
        "logs": logs,
        "detected_title": detected_title,
        "saved_path": saved_path,
    }


def default_output_path(input_path):
    root, ext = os.path.splitext(input_path)
    suffix = ext or ".docx"
    return f"{root}_Formatted{suffix}"


def parse_args(argv=None):
    ap = argparse.ArgumentParser(description="HNU Thesis Formatter v3")
    ap.add_argument("input", nargs="?", default=os.path.join(HERE, "thesis.docx"))
    ap.add_argument("output", nargs="?", default=None)
    ap.add_argument(
        "--title",
        default=None,
        help="Thesis title for even-page header (auto-detected if omitted)",
    )
    return ap.parse_args(argv)


def main(argv=None):
    args = parse_args(argv)
    input_path = args.input
    output_path = args.output or default_output_path(input_path)

    print(f"Input : {input_path}")
    print(f"Output: {output_path}\n")

    try:
        result = run_formatter(input_path, title=args.title, output_path=output_path, log_func=print)
        print(f"\nSaved: {result['saved_path']}")
    except PermissionError:
        import tempfile

        temp_path = os.path.join(tempfile.gettempdir(), os.path.basename(output_path))
        result = run_formatter(input_path, title=args.title, output_path=temp_path, log_func=print)
        print("\nOutput file is open/locked. Saved to temp path instead:")
        print(f"  {temp_path}")
        print("  Close the file in Word, then copy it manually.")

    print("Next step: open in Word -> Ctrl+A -> F9")


if __name__ == "__main__":
    main()
